# import re
# import PyPDF2
# import docx
# from sklearn.feature_extraction.text import TfidfVectorizer
# from sklearn.metrics.pairwise import cosine_similarity
# from skills import SKILLS_DB

# # ---------------------------
# # TEXT EXTRACTION
# # ---------------------------

# def extract_text_from_pdf(file):
#     text = ""
#     pdf = PyPDF2.PdfReader(file)
#     for page in pdf.pages:
#         if page.extract_text():
#             text += page.extract_text()
#     return text

# def extract_text_from_docx(file):
#     doc = docx.Document(file)
#     return " ".join([para.text for para in doc.paragraphs])

# # ---------------------------
# # PREPROCESSING
# # ---------------------------

# def preprocess(text):
#     text = text.lower()
#     text = re.sub(r'[^a-zA-Z ]', ' ', text)
#     text = re.sub(r'\s+', ' ', text)
#     return text

# # ---------------------------
# # SKILL EXTRACTION
# # ---------------------------

# def extract_skills(text):
#     found = []
#     for skill in SKILLS_DB:
#         if skill in text:
#             found.append(skill)
#     return list(set(found))

# # ---------------------------
# # SIMILARITY SCORE
# # ---------------------------

# def match_resume_jd(resume, jd):
#     tfidf = TfidfVectorizer()
#     vectors = tfidf.fit_transform([resume, jd])
#     score = cosine_similarity(vectors[0:1], vectors[1:2])
#     return round(score[0][0] * 100, 2)

# # ---------------------------
# # ATS SCORE
# # ---------------------------

# def ats_score(similarity, skills_found):
#     skill_score = len(skills_found) * 3
#     final_score = similarity + skill_score
#     return min(100, round(final_score, 2))

# # ---------------------------
# # SUGGESTIONS
# # ---------------------------

# def generate_suggestions(resume_text, jd_text):
#     suggestions = []
    
#     for skill in SKILLS_DB:
#         if skill in jd_text and skill not in resume_text:
#             suggestions.append(f"Add missing skill: {skill}")
    
#     if "project" not in resume_text:
#         suggestions.append("Add project section")
        
#     if "experience" not in resume_text:
#         suggestions.append("Mention experience clearly")
    
#     return suggestions

import re
import PyPDF2
import docx
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from skills import SKILLS_DB

# -------- TEXT EXTRACTION --------

def extract_text(file):
    if file.name.endswith(".pdf"):
        pdf = PyPDF2.PdfReader(file)
        return " ".join([page.extract_text() or "" for page in pdf.pages])

    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        return " ".join([para.text for para in doc.paragraphs])

    return ""

# -------- PREPROCESS --------

def preprocess(text):
    text = text.lower()
    text = re.sub(r'[^a-zA-Z ]', ' ', text)
    return re.sub(r'\s+', ' ', text)

# -------- SKILL EXTRACTION --------

def extract_skills(text):
    return list(set([skill for skill in SKILLS_DB if skill in text]))

# -------- SIMILARITY --------

def match_resume_jd(resume, jd):
    tfidf = TfidfVectorizer()
    vectors = tfidf.fit_transform([resume, jd])
    score = cosine_similarity(vectors[0:1], vectors[1:2])
    return round(score[0][0] * 100, 2)

# -------- ATS SCORE --------

def ats_score(similarity, skills_found):
    skill_score = len(skills_found) * 3
    return min(100, round(similarity + skill_score, 2))

# -------- SUGGESTIONS --------

def generate_suggestions(resume, jd):
    suggestions = []

    for skill in SKILLS_DB:
        if skill in jd and skill not in resume:
            suggestions.append(f"👉 Add missing skill: {skill}")

    if "project" not in resume:
        suggestions.append("👉 Add a PROJECT section")

    if "experience" not in resume:
        suggestions.append("👉 Add EXPERIENCE section")

    return suggestions